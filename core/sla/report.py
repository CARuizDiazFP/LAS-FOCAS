# Nombre de archivo: report.py
# Ubicación de archivo: core/sla/report.py
# Descripción: Render de informe SLA en DOCX/PDF
"""Generación de documentos DOCX/PDF para informes SLA."""

from __future__ import annotations

import copy
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.docx_utils.text_replace import replace_title_everywhere
from modules.common.libreoffice_export import convert_to_pdf

from .config import MESES_ES, REPORTS_DIR, SLA_TEMPLATE_PATH, SOFFICE_BIN
from .engine import SLAComputation, ServiceMetrics

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class DocumentoSLA:
    """Rutas de salida del informe."""

    docx: Path
    pdf: Optional[Path]


def generar_documento(
    resultado: SLAComputation,
    *,
    eventos: str = "",
    conclusion: str = "",
    propuesta: str = "",
    incluir_pdf: bool = False,
    reports_dir: Path | None = None,
    soffice_bin: Optional[str] = None,
) -> DocumentoSLA:
    """Crea el DOCX (y PDF opcional) usando la plantilla oficial."""

    if not SLA_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Plantilla SLA no encontrada en {SLA_TEMPLATE_PATH}")

    documento = _cargar_plantilla()
    mes_nombre = MESES_ES[resultado.mes - 1].capitalize()
    titulo = f"Informe SLA — {mes_nombre} {resultado.anio}"
    replace_title_everywhere(documento, titulo)

    _poblar_plantilla(
        documento,
        resultado,
        eventos=eventos,
        conclusion=conclusion,
        propuesta=propuesta,
    )

    destino_base = Path(reports_dir or REPORTS_DIR)
    destino = destino_base / "sla" / f"{resultado.anio:04d}{resultado.mes:02d}"
    destino.mkdir(parents=True, exist_ok=True)

    docx_path = destino / f"sla_{resultado.anio:04d}{resultado.mes:02d}.docx"
    documento.save(docx_path)

    pdf_path: Optional[Path] = None
    if incluir_pdf:
        binario = soffice_bin or SOFFICE_BIN
        if binario:
            try:
                pdf_generado = convert_to_pdf(str(docx_path), binario)
            except Exception:  # pragma: no cover - logging y fallback
                logger.exception("action=sla_report stage=pdf error")
            else:
                pdf_path = Path(pdf_generado)
        else:
            logger.info("action=sla_report stage=pdf reason=missing_soffice")

    return DocumentoSLA(docx=docx_path, pdf=pdf_path)


def _cargar_plantilla() -> Document:
    try:
        return Document(str(SLA_TEMPLATE_PATH))
    except Exception as exc:  # pragma: no cover - logging
        logger.exception("action=sla_report stage=load_template error=%s", exc)
        raise


def _poblar_plantilla(
    doc: Document,
    resultado: SLAComputation,
    *,
    eventos: str,
    conclusion: str,
    propuesta: str,
) -> None:
    if not doc.tables or len(doc.tables) < 3:
        raise ValueError("La plantilla debe incluir al menos tres tablas (resumen y bloques por servicio)")

    cuerpo = doc._body._element
    tabla_principal = doc.tables[0]
    tabla2_tpl = copy.deepcopy(doc.tables[1]._tbl)
    tabla3_tpl = copy.deepcopy(doc.tables[2]._tbl)

    idx_t2 = cuerpo.index(doc.tables[1]._tbl)
    idx_t3 = cuerpo.index(doc.tables[2]._tbl)

    parrafos_tpl: List[tuple[str, Optional[str]]] = []
    for elem in list(cuerpo[idx_t2 + 1 : idx_t3]):
        if elem.tag.endswith("p"):
            parrafo = Paragraph(elem, doc)
            parrafos_tpl.append((parrafo.text, parrafo.style.name if parrafo.style else None))
        cuerpo.remove(elem)

    while len(tabla_principal.rows) > 1:
        tabla_principal._tbl.remove(tabla_principal.rows[1]._tr)

    cuerpo.remove(doc.tables[2]._tbl)
    cuerpo.remove(doc.tables[1]._tbl)

    servicios_ordenados = sorted(
        resultado.servicios,
        key=lambda srv: (
            -(resultado.servicios_meta.get(srv.service_id or "", {}).get("sla_pct") or srv.disponibilidad_pct / 100.0),
            srv.service_id or "",
        ),
    )

    for metricas in servicios_ordenados:
        _agregar_fila_resumen(tabla_principal, metricas, resultado.servicios_meta)

    total_servicios = len(servicios_ordenados)
    for indice, metricas in enumerate(servicios_ordenados):
        elem2 = copy.deepcopy(tabla2_tpl)
        cuerpo.append(elem2)
        tabla_servicio = doc.tables[-1]
        _completar_tabla_servicio(tabla_servicio, metricas, resultado.servicios_meta)

        insertar_idx = cuerpo.index(elem2)
        _insertar_parrafos(doc, cuerpo, insertar_idx, parrafos_tpl, eventos, conclusion, propuesta)

        elem3 = copy.deepcopy(tabla3_tpl)
        cuerpo.insert(insertar_idx + len(parrafos_tpl) + 1, elem3)
        tabla_incidentes = doc.tables[-1]
        _completar_tabla_incidentes(tabla_incidentes, metricas)

        if indice < total_servicios - 1:
            salto = doc.add_page_break()
            cuerpo.remove(salto._p)
            cuerpo.insert(cuerpo.index(elem3) + 1, salto._p)


def _agregar_fila_resumen(
    tabla: Table,
    metricas: ServiceMetrics,
    servicios_meta: dict,
) -> None:
    fila_base = copy.deepcopy(tabla.rows[0]._tr)
    tabla._tbl.append(fila_base)
    celdas = tabla.rows[-1].cells

    meta = servicios_meta.get(metricas.service_id or "", {})

    celdas[0].text = (metricas.tipo_servicio or meta.get("tipo_servicio") or "-").strip()
    celdas[1].text = (metricas.service_id or "-").strip()
    celdas[2].text = (metricas.cliente or meta.get("cliente") or "-").strip()
    celdas[3].text = _fmt_timedelta(metricas.downtime_h)

    sla_obj = meta.get("sla_pct")
    if sla_obj is not None:
        celdas[4].text = f"{float(sla_obj) * 100:.2f}%"
    else:
        celdas[4].text = f"{metricas.disponibilidad_pct:.2f}%"


def _completar_tabla_servicio(
    tabla: Table,
    metricas: ServiceMetrics,
    servicios_meta: dict,
) -> None:
    meta = servicios_meta.get(metricas.service_id or "", {})

    valores = {
        "servicio": _combinar_textos(metricas.tipo_servicio, metricas.service_id),
        "cliente": metricas.cliente or meta.get("cliente") or "",
        "sla": _formatear_sla(meta, metricas),
    "ticket": ", ".join(_tickets_metricas(metricas)),
    "reclamo": ", ".join(_tickets_metricas(metricas)),
    }

    for fila in tabla.rows:
        for idx, celda in enumerate(fila.cells):
            contenido = celda.text.lower()
            for clave, valor in valores.items():
                if clave in contenido:
                    if len(fila.cells) > idx + 1 and not fila.cells[idx + 1].text.strip():
                        fila.cells[idx + 1].text = valor
                    else:
                        base = celda.text.split(":")[0].strip(": ")
                        fila.cells[idx].text = f"{base}: {valor}"
                    break


def _insertar_parrafos(
    doc: Document,
    cuerpo,
    idx: int,
    parrafos_tpl: Iterable[tuple[str, Optional[str]]],
    eventos: str,
    conclusion: str,
    propuesta: str,
) -> None:
    contenidos = list(parrafos_tpl) or [
        ("Eventos sucedidos de mayor impacto en SLA:", None),
        ("Conclusión:", None),
        ("Propuesta de mejora:", None),
    ]

    for base, estilo in contenidos:
        texto = base
        llave = base.lower()
        if "eventos" in llave:
            texto = base.rstrip() + (f" {eventos.strip()}" if eventos else "")
        elif "conclus" in llave:
            texto = base.rstrip() + (f" {conclusion.strip()}" if conclusion else "")
        elif "propuesta" in llave:
            texto = base.rstrip() + (f" {propuesta.strip()}" if propuesta else "")

        parrafo = doc.add_paragraph(texto, style=estilo)
        cuerpo.remove(parrafo._p)
        cuerpo.insert(idx + 1, parrafo._p)
        idx += 1


def _completar_tabla_incidentes(tabla: Table, metricas: ServiceMetrics) -> None:
    while len(tabla.rows) > 1:
        tabla._tbl.remove(tabla.rows[1]._tr)

    total_h = 0.0
    for intervalo in metricas.intervals:
        for incidente in intervalo.incidentes:
            fila = tabla.add_row().cells
            fila[0].text = metricas.service_id or "-"
            fila[1].text = incidente.ticket_id or "-"
            fila[2].text = f"{incidente.duracion_h:.2f}" if incidente.duracion_h is not None else ""
            fila[3].text = incidente.causal or ""
            fila[4].text = _formatear_fecha(incidente.inicio)

            if incidente.duracion_h:
                total_h += float(incidente.duracion_h)

    if total_h:
        totales = tabla.add_row().cells
        totales[0].text = "Total"
        totales[2].text = f"{total_h:.2f}"


def _tickets_metricas(metricas: ServiceMetrics) -> List[str]:
    encontrados: List[str] = []
    for intervalo in metricas.intervals:
        for ticket in intervalo.incident_ids:
            if ticket and ticket not in encontrados:
                encontrados.append(ticket)
    return encontrados


def _fmt_timedelta(horas: float) -> str:
    total_segundos = int(round(horas * 3600))
    horas_int, resto = divmod(total_segundos, 3600)
    minutos, segundos = divmod(resto, 60)
    return f"{horas_int:03d}:{minutos:02d}:{segundos:02d}"


def _formatear_fecha(valor) -> str:
    if valor is None or valor is pd.NA:
        return ""
    try:
        ts = pd.to_datetime(valor)
    except Exception:
        return str(valor)
    if pd.isna(ts):
        return ""
    meses = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]
    return f"{ts.day:02d}-{meses[ts.month - 1]}-{str(ts.year)[2:]}"


def _formatear_sla(meta: dict, metricas: ServiceMetrics) -> str:
    sla = meta.get("sla_pct")
    if sla is not None:
        return f"{float(sla) * 100:.2f}%"
    return f"Disponibilidad: {metricas.disponibilidad_pct:.2f}%"


def _combinar_textos(tipo: Optional[str], service_id: Optional[str]) -> str:
    partes = [p for p in [tipo, service_id] if p]
    return " ".join(partes)
