# Nombre de archivo: legacy_report.py
# Ubicación de archivo: core/sla/legacy_report.py
# Descripción: Generación del informe SLA replicando el comportamiento heredado
"""Render del informe SLA siguiendo el flujo legacy (Sandy)."""

from __future__ import annotations

import copy
import io
import logging
import unicodedata
from dataclasses import dataclass
from datetime import datetime, time, timedelta
from pathlib import Path
from typing import Dict, Iterable, Optional, Sequence

import pandas as pd
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.docx_utils.text_replace import replace_text_everywhere
from core.sla.config import MESES_ES, REPORTS_DIR, SLA_TEMPLATE_PATH, SOFFICE_BIN
from core.sla.report import DocumentoSLA
from modules.common.libreoffice_export import convert_to_pdf

logger = logging.getLogger(__name__)

EXCEL_EPOCH = datetime(1899, 12, 31)


def _copy_cell_borders(source_cell, target_cell) -> None:
    """Copia los bordes de una celda a otra preservando el formato."""
    try:
        source_tc = source_cell._element
        target_tc = target_cell._element
        
        # Obtener tcPr (table cell properties) de la celda origen
        source_tcPr = source_tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        if source_tcPr is None:
            return
        
        # Buscar tcBorders en la celda origen
        source_borders = source_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
        if source_borders is None:
            return
        
        # Obtener o crear tcPr en la celda destino
        target_tcPr = target_tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        if target_tcPr is None:
            target_tcPr = OxmlElement('w:tcPr')
            target_tc.insert(0, target_tcPr)
        
        # Eliminar bordes existentes en destino
        existing_borders = target_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
        if existing_borders is not None:
            target_tcPr.remove(existing_borders)
        
        # Copiar los bordes
        target_tcPr.append(copy.deepcopy(source_borders))
    except Exception as exc:  # noqa: BLE001
        logger.debug("action=copy_cell_borders error=%s", exc)


def _set_cell_text(cell, new_text: str) -> None:
    """Reemplaza el texto de una celda preservando su formato original."""
    if not cell.paragraphs:
        cell.text = new_text
        return
    
    # Preservar formato del primer párrafo y su primer run
    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        cell.text = new_text
        return
    
    # Mantener formato del primer run
    first_run = paragraph.runs[0]
    run_format = {
        'bold': first_run.bold,
        'italic': first_run.italic,
        'underline': first_run.underline,
        'font_name': first_run.font.name if first_run.font else None,
        'font_size': first_run.font.size if first_run.font else None,
        'font_color': first_run.font.color.rgb if first_run.font and first_run.font.color else None,
    }
    
    # Limpiar todos los runs del párrafo
    for run in paragraph.runs:
        run.text = ""
    
    # Limpiar párrafos adicionales
    for i in range(len(cell.paragraphs) - 1, 0, -1):
        p = cell.paragraphs[i]
        p._element.getparent().remove(p._element)
    
    # Crear nuevo run con el texto y formato preservado
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = new_text
    if run_format['bold'] is not None:
        run.bold = run_format['bold']
    if run_format['italic'] is not None:
        run.italic = run_format['italic']
    if run_format['underline'] is not None:
        run.underline = run_format['underline']
    if run.font and run_format['font_name']:
        run.font.name = run_format['font_name']
    if run.font and run_format['font_size']:
        run.font.size = run_format['font_size']
    if run.font and run_format['font_color']:
        run.font.color.rgb = run_format['font_color']


def _normalize(text: str) -> str:
    cleaned = unicodedata.normalize("NFKD", str(text)).encode("ASCII", "ignore").decode()
    cleaned = cleaned.replace("/", " ").replace("_", " ")
    return " ".join(cleaned.lower().strip().split())


SERVICIOS_REQUIRED: Dict[str, Sequence[str]] = {
    "tipo_servicio": ("tipo servicio",),
    "numero_linea": ("numero linea", "numero servicio", "numero primer servicio"),
    "nombre_cliente": ("nombre cliente", "cliente"),
    "horas_total": ("horas reclamos todos", "horas reclamos"),
    "sla": ("sla", "sla entregado", "% sla", "disponibilidad", "% disponibilidad"),
}

SERVICIOS_OPTIONAL: Dict[str, Sequence[str]] = {
    "domicilio": ("direccion servicio", "domicilio", "direccion"),
    "numero_primer_servicio": ("numero primer servicio",),
}

RECLAMOS_REQUIRED: Dict[str, Sequence[str]] = {
    "numero_linea": ("numero linea", "numero de linea", "numero linea reclamo", "numero primer servicio"),
    "ticket": ("numero reclamo", "n° reclamo", "n° de ticket", "numero ticket", "ticket"),
    # IMPORTANTE: Usar "Horas Netas Cierre Problema Reclamo" (columna P)
    "horas": ("horas netas cierre problema reclamo",),
    "tipo_solucion": ("tipo solucion reclamo", "tipo solución reclamo", "tipo solucion", "causa", "tipo"),
    "fecha_inicio": ("fecha inicio reclamo", "fecha inicio problema reclamo", "inicio"),
}

RECLAMOS_OPTIONAL: Dict[str, Sequence[str]] = {
    "fecha_cierre": ("fecha cierre reclamo", "fecha cierre problema reclamo", "cierre"),
    "descripcion": ("descripcion", "descripcion solucion reclamo", "descripción solución reclamo"),
    "numero_primer_servicio": ("numero primer servicio",),
    "horas_netas_cierre": (
        "horas netas cierre problema reclamo",
        "horas netas cierre",
        "horas netas problema cierre",
        "netas cierre problema",
        "horas cierre netas",
    ),
    "horas_totales_cierre": (
        "horas totales cierre problema reclamo",
        "horas totales cierre",
        "horas cierre problema",
        "total horas cierre",
    ),
}


@dataclass(slots=True)
class _ExcelDataset:
    dataframe: pd.DataFrame
    columns: Dict[str, str]
    optional: Dict[str, Optional[str]]


def _match_headers(
    columns: Iterable[str],
    required: Dict[str, Sequence[str]],
    optional: Dict[str, Sequence[str]] | None = None,
) -> tuple[Dict[str, str], Dict[str, Optional[str]]]:
    normalized = {_normalize(col): col for col in columns}

    def _find_match(candidates: Sequence[str]) -> Optional[str]:
        for candidate in candidates:
            normalized_candidate = _normalize(candidate)
            if normalized_candidate in normalized:
                return normalized[normalized_candidate]
        return None

    resolved: Dict[str, str] = {}
    for key, synonyms in required.items():
        found = _find_match(synonyms)
        if not found:
            raise ValueError(key)
        resolved[key] = found

    resolved_optional: Dict[str, Optional[str]] = {}
    if optional:
        for key, synonyms in optional.items():
            resolved_optional[key] = _find_match(synonyms)
    return resolved, resolved_optional


def _read_excel(content: bytes) -> pd.DataFrame:
    dataframe = pd.read_excel(io.BytesIO(content))
    dataframe.columns = dataframe.columns.astype(str).str.replace(r"\s+", " ", regex=True).str.strip()
    return dataframe


def _normalize_line_value(valor) -> str:
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return ""
    if isinstance(valor, float) and valor.is_integer():
        valor = int(valor)
    texto = str(valor).strip()
    if texto.endswith(".0") and texto.replace(".", "", 1).isdigit():
        texto = texto[:-2]
    return texto


def load_servicios_excel(content: bytes) -> _ExcelDataset:
    df = _read_excel(content)
    try:
        resolved, optional = _match_headers(df.columns, SERVICIOS_REQUIRED, SERVICIOS_OPTIONAL)
    except ValueError as exc:
        missing = str(exc)
        raise ValueError(f"Faltan columnas en Excel de servicios: {missing.upper()}") from exc

    df = df.copy()
    linea_col = resolved["numero_linea"]
    df[linea_col] = df[linea_col].apply(_normalize_line_value)
    numero_primer_servicio = optional.get("numero_primer_servicio")
    if numero_primer_servicio:
        df[numero_primer_servicio] = df[numero_primer_servicio].apply(_normalize_line_value)

    sla_col = resolved["sla"]
    df[sla_col] = df[sla_col].apply(_parse_sla_value)
    horas_col = resolved["horas_total"]
    df[horas_col] = df[horas_col].apply(_to_timedelta).apply(_fmt_td)
    df.sort_values(resolved["sla"], ascending=False, inplace=True)
    return _ExcelDataset(df.reset_index(drop=True), resolved, optional)


def load_reclamos_excel(content: bytes) -> _ExcelDataset:
    df = _read_excel(content)
    try:
        resolved, optional = _match_headers(df.columns, RECLAMOS_REQUIRED, RECLAMOS_OPTIONAL)
    except ValueError as exc:
        missing = str(exc)
        raise ValueError(f"Faltan columnas en Excel de reclamos: {missing.upper()}") from exc

    df = df.copy()
    horas_col = resolved["horas"]
    
    # DEBUG: Log valores crudos de la columna de horas antes de convertir
    logger.info(
        "action=sla_legacy_report stage=load_reclamos columna_horas=%s primeros_valores_crudos=%s tipos=%s",
        horas_col,
        df[horas_col].head(5).tolist(),
        [type(v).__name__ for v in df[horas_col].head(5).tolist()],
    )
    
    df[horas_col] = df[horas_col].apply(_horas_decimal)
    
    linea_col = resolved["numero_linea"]
    df[linea_col] = df[linea_col].apply(_normalize_line_value)
    numero_primer_servicio = optional.get("numero_primer_servicio")
    if numero_primer_servicio:
        df[numero_primer_servicio] = df[numero_primer_servicio].apply(_normalize_line_value)

    horas_netas_cierre_col = optional.get("horas_netas_cierre")
    if not horas_netas_cierre_col:
        raise ValueError(
            "Falta la columna 'Horas Netas Cierre Problema Reclamo' (columna P) en el Excel de reclamos"
        )

    logger.info(
        "action=sla_legacy_report stage=load_reclamos horas_netas_cierre_detectada=%s",
        horas_netas_cierre_col,
    )
    df[horas_netas_cierre_col] = df[horas_netas_cierre_col].apply(_horas_decimal)

    # Procesar columna de horas totales cierre si existe
    horas_totales_col = optional.get("horas_totales_cierre")
    if horas_totales_col:
        logger.info(
            "action=sla_legacy_report stage=load_reclamos horas_totales_cierre_detectada=%s",
            horas_totales_col
        )
        df[horas_totales_col] = df[horas_totales_col].apply(_horas_decimal)
    else:
        logger.warning("action=sla_legacy_report stage=load_reclamos horas_totales_cierre=no_detectada")
    
    fecha_col = resolved["fecha_inicio"]
    df[fecha_col] = df[fecha_col].apply(_to_datetime)
    cierre_col = optional.get("fecha_cierre")
    if cierre_col:
        df[cierre_col] = df[cierre_col].apply(_to_datetime)
    df.sort_values([resolved["numero_linea"], fecha_col], inplace=True)
    return _ExcelDataset(df.reset_index(drop=True), resolved, optional)


def identificar_excel(content: bytes) -> str:
    df = _read_excel(content)
    normalized = {_normalize(col) for col in df.columns}
    is_servicios = all(any(candidate in normalized for candidate in group) for group in SERVICIOS_REQUIRED.values())
    is_reclamos = all(any(candidate in normalized for candidate in group) for group in RECLAMOS_REQUIRED.values())
    if is_servicios and not is_reclamos:
        return "servicios"
    if is_reclamos and not is_servicios:
        return "reclamos"
    if is_servicios and is_reclamos:
        # Prefer reclamos si contiene más columnas informativas propias de tickets
        reclamo_specific = RECLAMOS_REQUIRED["ticket"].union(RECLAMOS_REQUIRED["tipo_solucion"])
        if any(candidate in normalized for candidate in reclamo_specific):
            return "reclamos"
        return "servicios"
    raise ValueError("No se pudo identificar el Excel (servicios vs reclamos)")


def generate_from_excel_pair(
    servicios_content: bytes,
    reclamos_content: bytes,
    *,
    mes: int,
    anio: int,
    incluir_pdf: bool = False,
    reports_dir: Path | None = None,
    soffice_bin: Optional[str] = None,
    eventos: str = "",
    conclusion: str = "",
    propuesta: str = "",
) -> DocumentoSLA:
    servicios = load_servicios_excel(servicios_content)
    reclamos = load_reclamos_excel(reclamos_content)
    return _render_document(
        servicios,
        reclamos,
        mes=mes,
        anio=anio,
        incluir_pdf=incluir_pdf,
        reports_dir=reports_dir or REPORTS_DIR,
        soffice_bin=soffice_bin or SOFFICE_BIN,
        eventos=eventos,
        conclusion=conclusion,
        propuesta=propuesta,
    )


def _render_document(
    servicios: _ExcelDataset,
    reclamos: _ExcelDataset,
    *,
    mes: int,
    anio: int,
    incluir_pdf: bool,
    reports_dir: Path,
    soffice_bin: Optional[str],
    eventos: str,
    conclusion: str,
    propuesta: str,
) -> DocumentoSLA:
    if not SLA_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Plantilla SLA no encontrada: {SLA_TEMPLATE_PATH}")

    doc = Document(str(SLA_TEMPLATE_PATH))
    
    # Actualizar el texto flotante con el formato: "Informe SLA Octubre 2025"
    mes_nombre = MESES_ES[mes - 1].capitalize()
    
    # Reemplazar los placeholders del template:
    # - "XXXXX" -> nombre del mes + espacio (ej: "Octubre ")
    # - "2023" -> año (ej: "2025")
    replace_text_everywhere(doc, {
        "XXXXX": f"{mes_nombre} ",  # Agregar espacio después del mes
        "2023": str(anio),
    })

    if len(doc.tables) < 3:
        raise ValueError("La plantilla SLA debe contener al menos tres tablas base")

    cuerpo = doc._body._element
    tabla_principal = doc.tables[0]
    tabla_servicio_tpl = copy.deepcopy(doc.tables[1]._tbl)
    tabla_detalle_tpl = copy.deepcopy(doc.tables[2]._tbl)

    idx_tabla2 = cuerpo.index(doc.tables[1]._tbl)
    idx_tabla3 = cuerpo.index(doc.tables[2]._tbl)

    bloques_texto: list[tuple[str, Optional[str]]] = []
    for element in list(cuerpo[idx_tabla2 + 1 : idx_tabla3]):
        if element.tag.endswith("p"):
            parrafo = Paragraph(element, doc)
            bloques_texto.append((parrafo.text, parrafo.style.name if parrafo.style else None))
        cuerpo.remove(element)

    while len(tabla_principal.rows) > 1:
        tabla_principal._tbl.remove(tabla_principal.rows[1]._tr)

    cuerpo.remove(doc.tables[2]._tbl)
    cuerpo.remove(doc.tables[1]._tbl)

    _rellenar_tabla_principal(tabla_principal, servicios, reclamos)

    num_servicios = len(servicios.dataframe)
    for idx, srv_row in servicios.dataframe.iterrows():
        elem_servicio = copy.deepcopy(tabla_servicio_tpl)
        cuerpo.append(elem_servicio)
        tabla_servicio = doc.tables[-1]
        _completar_tabla_servicio(tabla_servicio, srv_row, servicios, reclamos)

        insert_idx = cuerpo.index(elem_servicio)
        _insertar_bloques(doc, cuerpo, insert_idx, bloques_texto, eventos, conclusion, propuesta)

        elem_detalle = copy.deepcopy(tabla_detalle_tpl)
        cuerpo.insert(insert_idx + len(bloques_texto) + 1, elem_detalle)
        tabla_detalle = doc.tables[-1]
        _completar_tabla_detalle(tabla_detalle, srv_row, servicios, reclamos)

        if idx < num_servicios - 1:
            salto = doc.add_page_break()
            cuerpo.remove(salto._p)
            cuerpo.insert(cuerpo.index(elem_detalle) + 1, salto._p)

    destino = Path(reports_dir) / "sla" / f"{anio:04d}{mes:02d}"
    destino.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = destino / f"InformeSLA_{timestamp}.docx"
    doc.save(docx_path)

    pdf_path: Optional[Path] = None
    if incluir_pdf:
        binario = soffice_bin
        if binario:
            try:
                pdf_generado = convert_to_pdf(str(docx_path), binario)
                pdf_path = Path(pdf_generado)
            except Exception as exc:  # pragma: no cover - logging
                logger.warning("action=sla_legacy_report stage=pdf error=%s", exc)
        else:
            logger.info("action=sla_legacy_report stage=pdf reason=missing_soffice")

    return DocumentoSLA(docx=docx_path, pdf=pdf_path)


def _columna_horas_reclamos(reclamos: _ExcelDataset) -> tuple[str, str]:
    """Devuelve la columna de horas preferida para los cálculos de reclamos.
    
    IMPORTANTE: Usa la columna 'Horas Netas Cierre Problema Reclamo' (columna P del Excel),
    que es la columna principal de horas según requerimiento.
    """
    # Usar la columna 'horas' que mapea a 'Horas Netas Cierre Problema Reclamo' (columna P)
    columna = reclamos.columns.get("horas")
    if columna and columna in reclamos.dataframe.columns:
        return columna, "horas"
    raise ValueError(
        "Falta la columna 'Horas Netas Cierre Problema Reclamo' (columna P) en el Excel de reclamos"
    )


def _subset_reclamos_por_servicio(
    srv_row: pd.Series,
    servicios: _ExcelDataset,
    reclamos: _ExcelDataset,
) -> tuple[pd.DataFrame, str]:
    """Obtiene los reclamos asociados a un servicio considerando múltiples identificadores."""

    servicio_candidates: list[str] = []
    srv_linea_col = servicios.columns["numero_linea"]
    linea_val = str(srv_row.get(srv_linea_col, ""))
    if linea_val:
        servicio_candidates.append(linea_val)

    srv_linea_alt = servicios.optional.get("numero_primer_servicio")
    if srv_linea_alt:
        alt_val = str(srv_row.get(srv_linea_alt, ""))
        if alt_val and alt_val not in servicio_candidates:
            servicio_candidates.append(alt_val)

    recl_linea_col = reclamos.columns["numero_linea"]
    recl_alt_col = reclamos.optional.get("numero_primer_servicio")

    for candidate in servicio_candidates:
        if not candidate:
            continue
        mask = reclamos.dataframe[recl_linea_col] == candidate
        if recl_alt_col:
            mask = mask | (reclamos.dataframe[recl_alt_col] == candidate)
        subset = reclamos.dataframe[mask]
        if not subset.empty:
            linea_display = subset[recl_linea_col].dropna().iloc[0]
            return subset, str(linea_display)

    return pd.DataFrame(columns=reclamos.dataframe.columns), servicio_candidates[0] if servicio_candidates else ""


def _rellenar_tabla_principal(tabla: Table, servicios: _ExcelDataset, reclamos: _ExcelDataset) -> None:
    tipo_col = servicios.columns["tipo_servicio"]
    linea_col = servicios.columns["numero_linea"]
    cliente_col = servicios.columns["nombre_cliente"]
    sla_col = servicios.columns["sla"]
    
    recl_linea_col = reclamos.columns["numero_linea"]
    recl_horas, origen_horas = _columna_horas_reclamos(reclamos)
    logger.info(
        "action=sla_legacy_report stage=tabla_principal columna_horas=%s origen=%s",
        recl_horas,
        origen_horas,
    )

    while len(tabla.rows) > 1:
        tabla._tbl.remove(tabla.rows[1]._tr)

    # Guardar referencia a la fila template con sus bordes
    fila_template = tabla.rows[0]

    for _, fila in servicios.dataframe.iterrows():
        nueva = copy.deepcopy(fila_template._tr)
        tabla._tbl.append(nueva)
        celdas = tabla.rows[-1].cells
        
        # Copiar bordes de la fila template a las nuevas celdas
        for i, celda in enumerate(celdas):
            if i < len(fila_template.cells):
                _copy_cell_borders(fila_template.cells[i], celda)
        
        _set_cell_text(celdas[0], str(fila.get(tipo_col, "")))
        subset, linea_presentable = _subset_reclamos_por_servicio(fila, servicios, reclamos)
        _set_cell_text(celdas[1], linea_presentable or str(fila.get(linea_col, "")))
        _set_cell_text(celdas[2], str(fila.get(cliente_col, "")))
        
        if subset.empty:
            logger.warning(
                "action=sla_legacy_report stage=tabla_principal subset_vacio linea=%s",
                fila.get(linea_col),
            )

        total_horas = 0.0
        for _, reclamo in subset.iterrows():
            horas_val = reclamo.get(recl_horas)
            if not pd.isna(horas_val):
                total_horas += float(horas_val)

        if not subset.empty and total_horas <= 0:
            raise ValueError(
                f"No se pudieron sumar horas para el servicio {linea_presentable}. Verificar columna P del Excel de reclamos."
            )
        
        # Formatear como HHH:MM:SS
        horas_formateadas = _formatear_horas_totales(total_horas)
        _set_cell_text(celdas[3], horas_formateadas)
        
        sla_val = fila.get(sla_col)
        if pd.isna(sla_val):
            _set_cell_text(celdas[4], "")
        else:
            _set_cell_text(celdas[4], f"{float(sla_val) * 100:.2f}%")


def _completar_tabla_servicio(
    tabla: Table,
    srv_row: pd.Series,
    servicios: _ExcelDataset,
    reclamos: _ExcelDataset,
) -> None:
    valores = {
        "servicio": f"{srv_row.get(servicios.columns['tipo_servicio'], '')} {srv_row.get(servicios.columns['numero_linea'], '')}".strip(),
        "cliente": str(srv_row.get(servicios.columns['nombre_cliente'], "")),
        "ticket": ", ".join(_tickets_por_servicio(srv_row, servicios, reclamos)),
        "sla": _sla_texto(srv_row.get(servicios.columns['sla'])),
        "domicilio": _valor_opcional(srv_row, servicios.optional.get("domicilio")),
    }

    for fila in tabla.rows:
        for idx, celda in enumerate(fila.cells):
            contenido = celda.text.lower()
            for clave, valor in valores.items():
                if clave in contenido:
                    if len(fila.cells) > idx + 1 and not fila.cells[idx + 1].text.strip():
                        # Formato: Label en celda actual, valor en siguiente celda
                        # Agregar negrita al label
                        label_celda = fila.cells[idx]
                        label_text = label_celda.text
                        label_celda.text = ""
                        for paragraph in label_celda.paragraphs:
                            run = paragraph.add_run(label_text)
                            run.bold = True
                        
                        # Valor normal en siguiente celda
                        fila.cells[idx + 1].text = valor
                    else:
                        # Formato: "Label: valor" en la misma celda
                        base = celda.text.split(":")[0].strip(": ")
                        texto_completo = f"{base}: {valor}" if valor else base
                        
                        # Limpiar contenido actual
                        celda.text = ""
                        
                        # Crear runs con formato diferenciado
                        for paragraph in celda.paragraphs:
                            # Label en negrita
                            run_label = paragraph.add_run(f"{base}:")
                            run_label.bold = True
                            
                            # Valor en texto normal
                            if valor:
                                run_valor = paragraph.add_run(f" {valor}")
                                run_valor.bold = False
                    break


def _insertar_bloques(
    doc: Document,
    cuerpo,
    idx: int,
    bloques: Sequence[tuple[str, Optional[str]]],
    eventos: str,
    conclusion: str,
    propuesta: str,
) -> None:
    if not bloques:
        bloques = [
            ("Eventos sucedidos de mayor impacto en SLA:", None),
            ("Conclusión:", None),
            ("Propuesta de mejora:", None),
        ]
    for base, estilo in bloques:
        texto = base
        llave = base.lower()
        if "evento" in llave:
            texto = base.rstrip() + (" " + eventos.strip() if eventos else "")
        elif "conclusion" in llave or "conclusión" in llave:
            texto = base.rstrip() + (" " + conclusion.strip() if conclusion else "")
        elif "propuesta" in llave:
            texto = base.rstrip() + (" " + propuesta.strip() if propuesta else "")
        parrafo = doc.add_paragraph(texto, style=estilo)
        cuerpo.remove(parrafo._p)
        cuerpo.insert(idx + 1, parrafo._p)
        idx += 1


def _completar_tabla_detalle(
    tabla: Table,
    srv_row: pd.Series,
    servicios: _ExcelDataset,
    reclamos: _ExcelDataset,
) -> None:
    # Guardar fila template antes de eliminar filas
    fila_template_cells = [cell for cell in tabla.rows[0].cells] if len(tabla.rows) > 0 else []
    
    while len(tabla.rows) > 1:
        tabla._tbl.remove(tabla.rows[1]._tr)

    linea_col = servicios.columns["numero_linea"]
    recl_ticket = reclamos.columns["ticket"]

    recl_horas, _ = _columna_horas_reclamos(reclamos)

    recl_tipo = reclamos.columns["tipo_solucion"]
    recl_fecha = reclamos.columns["fecha_inicio"]
    recl_desc = reclamos.optional.get("descripcion")

    subset, _ = _subset_reclamos_por_servicio(srv_row, servicios, reclamos)

    total_horas = 0.0
    for _, reclamo in subset.iterrows():
        fila = tabla.add_row().cells
        
        # Copiar bordes de la fila template
        for i, celda in enumerate(fila):
            if i < len(fila_template_cells):
                _copy_cell_borders(fila_template_cells[i], celda)

        recl_linea_col = reclamos.columns["numero_linea"]
        _set_cell_text(fila[0], str(reclamo.get(recl_linea_col, "")))
        _set_cell_text(fila[1], str(reclamo.get(recl_ticket, "")))
        horas_val = reclamo.get(recl_horas)
        if pd.isna(horas_val):
            _set_cell_text(fila[2], "")
        else:
            _set_cell_text(fila[2], f"{horas_val:.2f}")
            total_horas += float(horas_val)
        _set_cell_text(fila[3], str(reclamo.get(recl_tipo, "")))
        fecha = reclamo.get(recl_fecha)
        _set_cell_text(fila[4], _formatear_fecha(fecha))
        if recl_desc and len(fila) > 5:
            _set_cell_text(fila[5], str(reclamo.get(recl_desc, "")))

    if total_horas:
        totales = tabla.add_row().cells
        
        # Copiar bordes de la fila template a la fila de totales
        for i, celda in enumerate(totales):
            if i < len(fila_template_cells):
                _copy_cell_borders(fila_template_cells[i], celda)
        
        _set_cell_text(totales[0], "Total")
        _set_cell_text(totales[2], f"{total_horas:.2f}")


def _tickets_por_servicio(
    srv_row: pd.Series,
    servicios: _ExcelDataset,
    reclamos: _ExcelDataset,
) -> list[str]:
    recl_ticket = reclamos.columns["ticket"]
    subset, _ = _subset_reclamos_por_servicio(srv_row, servicios, reclamos)
    tickets = subset[recl_ticket] if not subset.empty else pd.Series([], dtype=str)
    return [str(ticket) for ticket in tickets.dropna().unique()]


def _valor_opcional(row: pd.Series, column: Optional[str]) -> str:
    if not column:
        return ""
    return str(row.get(column, ""))


def _sla_texto(valor) -> str:
    if valor is None or pd.isna(valor):
        return ""
    try:
        return f"{float(valor) * 100:.2f}%"
    except Exception:  # noqa: BLE001
        return str(valor)


def _horas_decimal(valor) -> Optional[float]:
    """Convierte valores de horas a formato decimal.
    
    Maneja múltiples formatos de Excel:
    - datetime (Excel guarda tiempos > 24h como fechas desde 1899-12-31)
    - time (para tiempos < 24h)
    - timedelta
    - string formato HH:MM:SS
    - número decimal
    """
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return None

    if isinstance(valor, pd.Timestamp):
        valor = valor.to_pydatetime()

    # Excel guarda tiempos > 24h como datetime relativo a EXCEL_EPOCH (1899-12-31)
    if isinstance(valor, datetime):
        delta = valor - EXCEL_EPOCH
        return round(delta.total_seconds() / 3600, 4)

    if isinstance(valor, time):
        # Para tiempos < 24h que vienen como time object
        return round(valor.hour + valor.minute / 60 + valor.second / 3600, 4)

    if isinstance(valor, (pd.Timedelta, timedelta)):
        return round(valor.total_seconds() / 3600, 4)

    texto = str(valor).strip().lower().replace(",", ".")
    if not texto:
        return None
    try:
        if ":" in texto or "h" in texto:
            return round(float(pd.to_timedelta(texto).total_seconds() / 3600), 4)
        return round(float(texto), 4)
    except Exception:  # noqa: BLE001
        return None


def _to_datetime(valor):
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return None
    try:
        convertido = pd.to_datetime(valor, errors="coerce")
    except Exception:  # noqa: BLE001
        return None
    if pd.isna(convertido):
        return None
    return convertido


def _to_timedelta(valor) -> pd.Timedelta:
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return pd.Timedelta(0)
    texto = str(valor).strip().lower().replace(",", ".")
    if not texto:
        return pd.Timedelta(0)
    try:
        return pd.to_timedelta(texto)
    except Exception:  # noqa: BLE001
        try:
            return pd.to_timedelta(float(texto), unit="h")
        except Exception:  # noqa: BLE001
            return pd.Timedelta(0)


def _fmt_td(td: pd.Timedelta) -> str:
    total = int(td.total_seconds())
    horas, resto = divmod(total, 3600)
    minutos, segundos = divmod(resto, 60)
    return f"{horas:03d}:{minutos:02d}:{segundos:02d}"


def _formatear_horas_totales(horas_decimal: float) -> str:
    """Convierte horas decimales a formato HHH:MM:SS."""
    total_segundos = int(horas_decimal * 3600)
    horas, resto = divmod(total_segundos, 3600)
    minutos, segundos = divmod(resto, 60)
    return f"{horas:03d}:{minutos:02d}:{segundos:02d}"


def _formatear_fecha(valor) -> str:
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return ""
    try:
        fecha = pd.to_datetime(valor)
    except Exception:  # noqa: BLE001
        return str(valor)
    if pd.isna(fecha):
        return ""
    meses = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]
    return f"{fecha.day:02d}-{meses[fecha.month - 1]}-{str(fecha.year)[2:]}"


def _parse_sla_value(valor) -> Optional[float]:
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return None
    texto = str(valor).strip().lower().replace(",", ".")
    if not texto:
        return None
    try:
        numero = float(texto)
        if numero > 1:
            numero /= 100
        return numero
    except Exception:  # noqa: BLE001
        return None
