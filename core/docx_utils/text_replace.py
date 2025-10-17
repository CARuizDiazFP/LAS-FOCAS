# Nombre de archivo: text_replace.py
# Ubicación de archivo: core/docx_utils/text_replace.py
# Descripción: Utilidades para reemplazar texto en archivos DOCX, incluyendo shapes y encabezados

from __future__ import annotations

from typing import Dict, Iterable, List, Sequence, Set

from docx import Document
from docx.opc.part import Part


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

TITLE_PLACEHOLDER = "Informe Repetitividad Mes Año"


def replace_text_everywhere(document: Document, mapping: Dict[str, str]) -> int:
    """Reemplaza texto en todo el árbol XML del documento.

    La búsqueda cubre:
    - Texto regular (`w:t`).
    - Texto dentro de cuadros (`w:txbxContent//w:t`).
    - Texto contenido en dibujos (`a:t`).

    Args:
        document: Documento `python-docx` sobre el cual operar.
        mapping: Diccionario con pares ``texto_original`` → ``texto_nuevo``.

    Returns:
        Número de nodos cuyo contenido fue modificado.
    """

    if not mapping:
        return 0

    replacements = 0

    for part in _iter_unique_parts(document):
        element = getattr(part, "element", None)
        if element is None:
            continue

        for old, new in mapping.items():
            replacements += _replace_runs(element, old, new)

    return replacements


def replace_title_everywhere(document: Document, titulo: str) -> int:
    """Reemplaza el título del informe en cualquier parte del DOCX.

    Si no se detecta el placeholder original, inserta un párrafo nuevo al inicio.
    """

    mapping = {TITLE_PLACEHOLDER: titulo}
    replacements = 0

    for part in _iter_unique_parts(document):
        element = getattr(part, "element", None)
        if element is None:
            continue
        replacements += _replace_runs(element, TITLE_PLACEHOLDER, titulo)

    if replacements == 0:
        _ensure_title_paragraph(document, titulo)
        replacements = 1

    return replacements


def _replace_runs(element, needle: str, replacement: str) -> int:
    replaced = 0

    w_nodes = _safe_xpath(element, f".//*[local-name()='t' and namespace-uri()='{W_NS}']")
    if w_nodes:
        replaced += _replace_split_runs(w_nodes, needle, replacement)

    a_nodes = _safe_xpath(element, f".//*[local-name()='t' and namespace-uri()='{A_NS}']")
    if a_nodes:
        replaced += _replace_split_runs(a_nodes, needle, replacement)

    return replaced


def _replace_split_runs(nodes: Sequence, needle: str, replacement: str) -> int:
    needle_len = len(needle)
    max_len = needle_len + 10
    count = 0
    index = 0
    total_nodes = len(nodes)

    while index < total_nodes:
        node = nodes[index]
        text = getattr(node, "text", None) or ""

        if needle in text:
            node.text = text.replace(needle, replacement)
            count += 1
            index += 1
            continue

        buffer = text
        consumed: List = [node]
        j = index + 1
        replaced_here = False

        while j < total_nodes and len(buffer) <= max_len:
            next_node = nodes[j]
            buffer += getattr(next_node, "text", None) or ""
            consumed.append(next_node)
            if needle in buffer:
                new_text = buffer.replace(needle, replacement, 1)
                consumed[0].text = new_text
                for extra in consumed[1:]:
                    extra.text = ""
                count += 1
                index = j + 1
                replaced_here = True
                break
            j += 1

        if not replaced_here:
            index += 1

    return count


def _ensure_title_paragraph(document: Document, titulo: str) -> None:
    if document.paragraphs:
        paragraph = document.paragraphs[0].insert_paragraph_before(titulo)
    else:
        paragraph = document.add_paragraph(titulo)

    try:
        paragraph.style = document.styles["Title"]
    except Exception:  # noqa: BLE001 - estilo no disponible en todas las plantillas
        try:
            paragraph.style = "Title"
        except Exception:  # pragma: no cover - fallback silencioso
            pass


def _safe_xpath(element, expression: str):
    try:
        return element.xpath(expression)
    except Exception:  # pragma: no cover - protección ante expresiones inconsistentes
        return []


def _iter_unique_parts(document: Document) -> Iterable[Part]:
    """Itera por todas las partes únicas del documento (cuerpo, encabezados, pies, etc.)."""

    stack = [document.part]
    visited_by_name: Set[str] = set()
    visited_by_id: Set[int] = set()

    while stack:
        part = stack.pop()
        partname = getattr(part, "partname", None)
        part_id = id(part)

        if partname is not None:
            if partname in visited_by_name:
                continue
            visited_by_name.add(partname)
        else:
            if part_id in visited_by_id:
                continue
            visited_by_id.add(part_id)

        yield part

        rels = getattr(part, "rels", {})
        for rel in rels.values():  # type: ignore[call-arg]
            target = getattr(rel, "target_part", None)
            if target is None:
                continue
            target_name = getattr(target, "partname", None)
            target_id = id(target)
            if target_name is not None and target_name in visited_by_name:
                continue
            if target_name is None and target_id in visited_by_id:
                continue
            stack.append(target)
