# Nombre de archivo: test_repetitividad_docx_render.py
# Ubicación de archivo: tests/test_repetitividad_docx_render.py
# Descripción: Pruebas del render DOCX para bloques de servicios en repetitividad

import base64

from docx import Document

from modules.informes_repetitividad.report import (
    TABLE_HEADERS,
    MAP_MAX_HEIGHT,
    MAP_MAX_WIDTH,
    _format_horas,
    _render_service_block,
)
from modules.informes_repetitividad.schemas import ReclamoDetalle, ServicioDetalle


def _create_dummy_png(path):
    png_bytes = base64.b64decode(
        b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8Xw8AAn8B9GoPzZwAAAAASUVORK5CYII="
    )
    path.write_bytes(png_bytes)


def test_render_service_block_removes_lat_lon_columns(tmp_path):
    documento = Document()
    servicio = ServicioDetalle(
        servicio="Linea-Test",
        nombre_cliente="Cliente 1",
        tipo_servicio="FTTH",
        casos=2,
        reclamos=[
            ReclamoDetalle(
                numero_reclamo="123",
                numero_evento="E-1",
                fecha_inicio="2024-07-01",
                fecha_cierre="2024-07-02",
                horas_netas=750,
                tipo_solucion="Reparación",
                descripcion_solucion="Cambio de ONT",
                latitud=-34.60,
                longitud=-58.38,
            ),
            ReclamoDetalle(
                numero_reclamo="456",
                numero_evento="E-2",
                fecha_inicio="2024-07-10",
                fecha_cierre="2024-07-11",
                horas_netas=300,
                tipo_solucion="Reconfiguración",
                descripcion_solucion="Ajuste de OLT",
                latitud=-34.61,
                longitud=-58.39,
            ),
        ],
    )

    png_path = tmp_path / "map.png"
    _create_dummy_png(png_path)
    servicio.map_image_path = str(png_path)

    assert len(documento.tables) == 0
    _render_service_block(documento, servicio, with_geo=True)

    assert len(documento.tables) == 1
    tabla = documento.tables[0]
    assert len(tabla.columns) == len(TABLE_HEADERS)
    headers_obtenidos = [cell.text for cell in tabla.rows[0].cells]
    assert headers_obtenidos == TABLE_HEADERS
    assert all("Lat" not in header for header in headers_obtenidos)

    celdas_datos = [
        cell.text
        for row in tabla.rows[1:]
        for cell in row.cells
    ]
    assert all("-58.38" not in celda for celda in celdas_datos)
    assert "12:30" in celdas_datos
    assert "5:00" in celdas_datos
    assert any("Mapa georreferenciado" in par.text for par in documento.paragraphs)
    assert len(documento.inline_shapes) == 1
    picture = documento.inline_shapes[0]
    assert float(picture.width) <= float(MAP_MAX_WIDTH)
    assert float(picture.height) <= float(MAP_MAX_HEIGHT)


def test_format_horas_handles_various_inputs():
    class DummyDelta:
        def total_seconds(self):
            return 5400

    assert _format_horas(125) == "2:05"
    assert _format_horas(300.0) == "5:00"
    assert _format_horas("01:30") == "1:30"
    assert _format_horas("1:15:30") == "1:15"
    assert _format_horas(DummyDelta()) == "1:30"
    assert _format_horas(None) == "-"
    assert _format_horas(" ") == "-"
