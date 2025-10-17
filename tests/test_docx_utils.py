# Nombre de archivo: test_docx_utils.py
# Ubicación de archivo: tests/test_docx_utils.py
# Descripción: Pruebas unitarias para utilidades de documentos DOCX

from docx import Document
from docx.oxml import parse_xml

from core.docx_utils.text_replace import replace_text_everywhere, replace_title_everywhere


def test_replace_title_everywhere_replaces_in_shapes_and_drawings():
    document = Document()
    document.add_paragraph("Informe Repetitividad Mes Año")

    header_paragraph = document.sections[0].header.paragraphs[0]
    header_paragraph.text = "Informe Repetitividad Mes Año"

    textbox_xml = parse_xml(
        "<w:p xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main' "
        "xmlns:v='urn:schemas-microsoft-com:vml'>"
        "<w:r><w:pict><v:shape id='TextBox1'><v:textbox><w:txbxContent>"
        "<w:p><w:r><w:t>Informe Repetitividad Mes Año</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    document.element.body.append(textbox_xml)

    drawing_xml = parse_xml(
        "<w:p xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main' "
        "xmlns:wp='http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing' "
        "xmlns:a='http://schemas.openxmlformats.org/drawingml/2006/main'>"
        "<w:r><w:drawing><wp:inline>"
        "<a:graphic><a:graphicData uri='http://schemas.openxmlformats.org/drawingml/2006/diagram'>"
        "<a:txBody><a:p><a:r><a:t>Informe Repetitividad Mes Año</a:t></a:r></a:p></a:txBody>"
        "</a:graphicData></a:graphic>"
        "</wp:inline></w:drawing></w:r></w:p>"
    )
    document.element.body.append(drawing_xml)

    nuevo_texto = "Informe Repetitividad — Octubre 2024"
    reemplazos = replace_title_everywhere(document, nuevo_texto)

    assert reemplazos >= 1
    assert document.paragraphs[0].text == nuevo_texto
    assert document.sections[0].header.paragraphs[0].text == nuevo_texto

    textbox_nodes = document.element.xpath(".//*[local-name()='txbxContent']//*[local-name()='t']")
    assert textbox_nodes and all(getattr(node, "text", None) == nuevo_texto for node in textbox_nodes)

    drawing_nodes = document.element.xpath(".//*[local-name()='drawing']//*[local-name()='t']")
    assert drawing_nodes and all(getattr(node, "text", None) == nuevo_texto for node in drawing_nodes)


def test_replace_title_everywhere_inserts_fallback_when_placeholder_missing():
    document = Document()
    document.add_paragraph("Sin placeholder")

    nuevo_texto = "Informe Repetitividad — Octubre 2024"
    reemplazos = replace_title_everywhere(document, nuevo_texto)

    assert reemplazos == 1
    assert document.paragraphs[0].text == nuevo_texto
    assert document.paragraphs[1].text == "Sin placeholder"


def test_replace_text_everywhere_general_mapping():
    document = Document()
    document.add_paragraph("Hola mundo")
    reemplazos = replace_text_everywhere(document, {"Hola": "Chau"})
    assert reemplazos == 1
    assert document.paragraphs[0].text == "Chau mundo"
