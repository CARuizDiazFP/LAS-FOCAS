# Nombre de archivo: test_report_builder.py
# Ubicación de archivo: tests/test_report_builder.py
# Descripción: Pruebas de generación de reportes DOCX

from pathlib import Path

from docx import Document

from modules.informes_repetitividad.report import export_docx
from modules.informes_repetitividad.schemas import (
    Params,
    ReclamoDetalle,
    ResultadoRepetitividad,
    ServicioDetalle,
)


def test_export_docx_crea_archivo(tmp_path):
    params = Params(periodo_mes=7, periodo_anio=2024)
    data = ResultadoRepetitividad(
        servicios=[
            ServicioDetalle(
                servicio="S1",
                nombre_cliente="Cliente Demo",
                tipo_servicio="Fibra",
                casos=2,
                reclamos=[
                    ReclamoDetalle(
                        numero_reclamo="1",
                        fecha_inicio="2024-07-01",
                        fecha_cierre="2024-07-02",
                        horas_netas=12.5,
                        descripcion_solucion="Cambio de tarjeta",
                    ),
                    ReclamoDetalle(
                        numero_reclamo="2",
                        numero_evento="EVT-77",
                        descripcion_solucion="Reinicio programado",
                    ),
                ],
            )
        ],
        total_servicios=4,
        total_repetitivos=1,
        periodos=["2024-07"],
        with_geo=False,
        source="excel",
    )
    path = export_docx(data, params, tmp_path, with_geo=False)
    assert Path(path).exists()
    doc = Document(path)
    assert any("Julio 2024" in p.text for p in doc.paragraphs)
    assert any("Cliente Demo" in p.text for p in doc.paragraphs)
    assert any("Fibra" in p.text for p in doc.paragraphs)
