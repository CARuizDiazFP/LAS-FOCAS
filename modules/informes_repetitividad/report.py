# Nombre de archivo: report.py
# Ubicación de archivo: modules/informes_repetitividad/report.py
# Descripción: Generación de archivos DOCX y PDF para el informe de repetitividad

import logging
from pathlib import Path
from typing import Any, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.shape import InlineShape

from modules.common.libreoffice_export import convert_to_pdf
from core.docx_utils.text_replace import replace_title_everywhere
from core.maps.static_map import build_static_map_png
from core.utils.timefmt import minutes_to_hhmm
from .config import (
    MAPS_ENABLED,
    MESES_ES,
    REP_TEMPLATE_PATH,
)
from .schemas import Params, ResultadoRepetitividad, ServicioDetalle, ReclamoDetalle

logger = logging.getLogger(__name__)

TABLE_HEADERS = [
    "N° Reclamo",
    "N° Evento",
    "Fecha Inicio",
    "Fecha Cierre",
    "Horas Netas",
    "Tipo Solución",
    "Descripción",
]

MAP_MAX_WIDTH = Inches(6.2)
MAP_MAX_HEIGHT = Inches(5.6)


def _header_cell(cell, text: str) -> None:
    cell.text = text
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "D9D9D9")
    cell._tc.get_or_add_tcPr().append(shading)


def _load_template() -> Document:
    """Carga la plantilla oficial o crea un documento vacío como fallback."""

    if REP_TEMPLATE_PATH.exists():
        try:
            return Document(str(REP_TEMPLATE_PATH))
        except Exception:  # pragma: no cover - logging
            logger.exception("action=load_template error path=%s", REP_TEMPLATE_PATH)
    else:
        logger.warning("Plantilla de repetitividad no encontrada en %s", REP_TEMPLATE_PATH)
    return Document()


def export_docx(
    data: ResultadoRepetitividad,
    periodo: Params,
    out_dir: str,
    with_geo: bool = False,
) -> str:
    """Genera el archivo DOCX con bloques por servicio y mapa opcional."""

    mes_nombre = MESES_ES[periodo.periodo_mes - 1].capitalize()
    doc = _load_template()

    titulo = f"Informe Repetitividad — {mes_nombre} {periodo.periodo_anio}"
    replace_title_everywhere(doc, titulo)

    # Título principal
    doc.add_heading(
        titulo,
        level=1,
    )

    # Resumen ejecutivo
    porcentaje = 100 * data.total_repetitivos / max(data.total_servicios, 1)
    doc.add_paragraph(
        f"Servicios analizados: {data.total_servicios} | "
        f"Servicios con repetitividad: {data.total_repetitivos} ({porcentaje:.1f}%)"
    )

    if data.periodos:
        doc.add_paragraph(
            "Períodos detectados: " + ", ".join(sorted(set(data.periodos)))
        )

    if not data.servicios:
        doc.add_paragraph("No se detectaron servicios repetitivos en el período seleccionado.")
    else:
        for servicio in data.servicios:
            _render_service_block(doc, servicio, with_geo)

    out_dir_path = Path(out_dir)
    out_dir_path.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir_path / f"repetitividad_{periodo.periodo_anio}{periodo.periodo_mes:02d}.docx"
    doc.save(docx_path)
    logger.info("action=export_docx path=%s servicios=%s", docx_path, len(data.servicios))
    return str(docx_path)


def _render_service_block(doc: Document, servicio: ServicioDetalle, with_geo: bool) -> None:
    """Dibuja el bloque detallado para un servicio repetitivo."""

    heading_parts: List[str] = []
    if servicio.tipo_servicio:
        heading_parts.append(servicio.tipo_servicio)
    heading_parts.append(servicio.servicio)
    heading = " — ".join(filter(None, heading_parts)) or servicio.servicio
    if servicio.nombre_cliente:
        heading = f"{heading} · {servicio.nombre_cliente}"
    doc.add_heading(heading, level=2)

    table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(TABLE_HEADERS):
        _header_cell(hdr_cells[idx], header)

    for reclamo in servicio.reclamos:
        _render_reclamo_row(table, reclamo)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if with_geo and servicio.map_image_path:
        image_path = Path(servicio.map_image_path)
        if image_path.exists():
            try:
                _insert_service_map(doc, image_path)
            except Exception as exc:  # noqa: BLE001
                logger.warning("action=servicio_block image_failed map=%s error=%s", image_path, exc)
                doc.add_paragraph(f"No se pudo insertar el mapa ({image_path.name}).")
        else:
            logger.debug(
                "action=servicio_block map_missing servicio=%s path=%s",
                servicio.servicio,
                image_path,
            )

    doc.add_paragraph()


def _render_reclamo_row(table, reclamo: ReclamoDetalle) -> None:
    row_cells = table.add_row().cells
    row_cells[0].text = _safe_text(reclamo.numero_reclamo)
    row_cells[1].text = _safe_text(reclamo.numero_evento)
    row_cells[2].text = _safe_text(reclamo.fecha_inicio)
    row_cells[3].text = _safe_text(reclamo.fecha_cierre)
    row_cells[4].text = _format_horas(reclamo.horas_netas)
    row_cells[5].text = _safe_text(reclamo.tipo_solucion)
    row_cells[6].text = _shorten(_safe_text(reclamo.descripcion_solucion))


def _safe_text(value: Optional[str]) -> str:
    if value is None:
        return "-"
    text = str(value).strip()
    return text or "-"


def _shorten(value: str, limit: int = 220) -> str:
    return value if len(value) <= limit else value[: limit - 1] + "…"


def _format_horas(value: Any) -> str:
    return minutes_to_hhmm(value)


def _insert_service_map(doc: Document, image_path: Path) -> None:
    doc.add_paragraph()
    doc.add_paragraph("Mapa georreferenciado:")
    image_paragraph = doc.add_paragraph()
    run = image_paragraph.add_run()
    picture = run.add_picture(str(image_path))
    _fit_inline_picture(picture)


def _fit_inline_picture(picture: InlineShape) -> None:
    width = float(picture.width)
    height = float(picture.height)

    if height > float(MAP_MAX_HEIGHT):
        ratio = float(MAP_MAX_HEIGHT) / height
        picture.height = int(MAP_MAX_HEIGHT)
        picture.width = int(width * ratio)
        width = float(picture.width)
        height = float(picture.height)

    if width > float(MAP_MAX_WIDTH):
        ratio = float(MAP_MAX_WIDTH) / width
        picture.width = int(MAP_MAX_WIDTH)
        picture.height = int(height * ratio)


def maybe_export_pdf(docx_path: str, soffice_bin: Optional[str]) -> Optional[str]:
    """Convierte el DOCX a PDF si LibreOffice está disponible."""

    if not soffice_bin:
        logger.debug("action=maybe_export_pdf reason=missing_binary")
        return None

    binary_path = Path(soffice_bin)
    if not binary_path.exists():
        logger.info(
            "action=maybe_export_pdf reason=binary_not_found soffice_bin=%s",
            soffice_bin,
        )
        return None

    try:
        return convert_to_pdf(docx_path, soffice_bin)
    except Exception:  # pragma: no cover - logging
        logger.exception("action=maybe_export_pdf error")
        return None


def generate_service_maps(
    data: ResultadoRepetitividad,
    periodo: Params,
    out_dir: str,
    with_geo: bool,
) -> List[Path]:
    """Genera mapas por servicio y adjunta la ruta al resultado."""

    if not MAPS_ENABLED or not with_geo or not data.servicios:
        return []

    out_dir_path = Path(out_dir)
    out_dir_path.mkdir(parents=True, exist_ok=True)
    maps_dir = out_dir_path / f"maps_{periodo.periodo_anio}{periodo.periodo_mes:02d}"
    maps_dir.mkdir(parents=True, exist_ok=True)

    created: List[Path] = []

    for servicio in data.servicios:
        coords = _collect_coords(servicio)
        servicio.map_path = None
        if not coords:
            servicio.map_path = None
            servicio.map_image_path = None
            continue

        safe_name = "".join(ch if ch.isalnum() else "_" for ch in servicio.servicio)[:50]
        png_path = maps_dir / f"repetitividad_{periodo.periodo_anio}{periodo.periodo_mes:02d}_{safe_name}.png"

        try:
            build_static_map_png([(lat, lon) for lat, lon, _ in coords], png_path)
        except ValueError:
            servicio.map_image_path = None
            logger.debug("action=generate_service_maps reason=no_valid_points servicio=%s", servicio.servicio)
            continue
        except Exception as exc:  # noqa: BLE001
            servicio.map_image_path = None
            logger.warning(
                "action=generate_service_maps stage=static_map_failed servicio=%s error=%s",
                servicio.servicio,
                exc,
            )
            continue

        servicio.map_image_path = str(png_path)
        created.append(png_path)

    logger.info("action=generate_service_maps total=%s", len(created))
    return created


def _collect_coords(servicio: ServicioDetalle) -> List[tuple[float, float, Optional[str]]]:
    coords: List[tuple[float, float, Optional[str]]] = []
    for reclamo in servicio.reclamos:
        lat = _to_float(reclamo.latitud)
        lon = _to_float(reclamo.longitud)
        if lat is None or lon is None:
            continue
        coords.append((lat, lon, reclamo.numero_reclamo))
    return coords


def _to_float(value: Optional[float]) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):  # noqa: BLE001
        return None


def generate_geo_map(
    data: ResultadoRepetitividad,
    periodo: Params,
    out_dir: str,
) -> Optional[str]:
    """Compatibilidad: genera el primer mapa disponible o retorna None."""

    maps = generate_service_maps(data, periodo, out_dir, with_geo=True)
    return str(maps[0]) if maps else None
