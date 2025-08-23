# Nombre de archivo: report.py
# Ubicación de archivo: modules/informes_sla/report.py
# Descripción: Generación de archivos DOCX y PDF para el informe de SLA

import logging
import subprocess
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from modules.common.libreoffice_export import convert_to_pdf
from .config import MESES_ES
from .schemas import Params, ResultadoSLA

logger = logging.getLogger(__name__)


def _header_cell(cell, text: str) -> None:
    cell.text = text
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "D9D9D9")
    cell._tc.get_or_add_tcPr().append(shading)


def export_docx(data: ResultadoSLA, periodo: Params, out_dir: str) -> str:
    """Genera el documento DOCX con el análisis de SLA."""
    mes_nombre = MESES_ES[periodo.periodo_mes - 1].capitalize()
    doc = Document()
    doc.add_heading(f"Análisis de SLA — {mes_nombre} {periodo.periodo_anio}", level=1)

    kpi = data.kpi
    doc.add_paragraph(
        " | ".join(
            [
                f"Total: {kpi.total}",
                f"Cumplidos: {kpi.cumplidos}",
                f"Incumplidos: {kpi.incumplidos}",
                f"% Cumplimiento: {kpi.pct_cumplimiento:.2f}",
                f"TTR Promedio (h): {kpi.ttr_promedio_h:.2f}",
                f"TTR Mediana (h): {kpi.ttr_mediana_h:.2f}",
            ]
        )
    )
    if data.sin_cierre:
        doc.add_paragraph(f"Casos sin cierre excluidos: {data.sin_cierre}")

    tabla = doc.add_table(rows=1, cols=6)
    hdr = tabla.rows[0].cells
    _header_cell(hdr[0], "ID")
    _header_cell(hdr[1], "Cliente")
    _header_cell(hdr[2], "Servicio")
    _header_cell(hdr[3], "TTR (h)")
    _header_cell(hdr[4], "SLA Obj (h)")
    _header_cell(hdr[5], "Cumplido")

    for item in data.detalle[:2000]:
        row = tabla.add_row().cells
        row[0].text = item.id
        row[1].text = item.cliente
        row[2].text = item.servicio
        row[3].text = f"{item.ttr_h:.2f}"
        row[4].text = f"{item.sla_objetivo_h:.2f}"
        row[5].text = "Sí" if item.cumplido else "No"
        if not item.cumplido:
            row[5].paragraphs[0].runs[0].bold = True
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    if data.breakdown_por_servicio:
        doc.add_paragraph("")
        tabla_srv = doc.add_table(rows=1, cols=7)
        hdr2 = tabla_srv.rows[0].cells
        _header_cell(hdr2[0], "Servicio")
        _header_cell(hdr2[1], "Total")
        _header_cell(hdr2[2], "Cumplidos")
        _header_cell(hdr2[3], "Incumplidos")
        _header_cell(hdr2[4], "% Cumpl.")
        _header_cell(hdr2[5], "TTR Prom.")
        _header_cell(hdr2[6], "TTR Med.")
        for servicio, kpi_s in data.breakdown_por_servicio.items():
            row = tabla_srv.add_row().cells
            row[0].text = servicio
            row[1].text = str(kpi_s.total)
            row[2].text = str(kpi_s.cumplidos)
            row[3].text = str(kpi_s.incumplidos)
            row[4].text = f"{kpi_s.pct_cumplimiento:.2f}"
            row[5].text = f"{kpi_s.ttr_promedio_h:.2f}"
            row[6].text = f"{kpi_s.ttr_mediana_h:.2f}"
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)
    docx_path = out_path / f"sla_{periodo.periodo_anio}{periodo.periodo_mes:02d}.docx"
    doc.save(docx_path)
    docx_path.chmod(0o600)
    logger.info("action=export_docx path=%s", docx_path)
    return str(docx_path)


def maybe_export_pdf(docx_path: str, soffice_bin: Optional[str]) -> Optional[str]:
    """Convierte el DOCX a PDF si LibreOffice está disponible."""
    if not soffice_bin:
        return None
    try:
        pdf_path = convert_to_pdf(docx_path, soffice_bin)
        Path(pdf_path).chmod(0o600)
        return pdf_path
    except (FileNotFoundError, subprocess.CalledProcessError):  # pragma: no cover - logging
        logger.exception("action=maybe_export_pdf error")
        return None
    except Exception:  # pragma: no cover - logging
        logger.exception("action=maybe_export_pdf error_desconocido")
        return None
